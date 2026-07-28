from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
MEDIA_DIR = ROOT / "assets" / "media"
TOOLS_DIR = ROOT / "assets" / "tools"
MEDIA_DIR.mkdir(parents=True, exist_ok=True)
TOOLS_DIR.mkdir(parents=True, exist_ok=True)

DARK = "192720"
INK = "17211D"
MUTED = "66716B"
ACCENT = "C5DB49"
ACCENT_DARK = "738125"
SOFT = "E5EBE4"
PALE = "F4F1EA"
WHITE = "FFFFFF"
LINE = "D7DDD8"


def set_run(run, *, font="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, end])


def configure_document(doc, running_label):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, DARK),
        ("Heading 2", 13, 14, 7, DARK),
        ("Heading 3", 12, 10, 5, ACCENT_DARK),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(running_label.upper()), size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    set_run(p.add_run("MICHEAL SHEEHY  ·  "), size=8.5, color=MUTED, bold=True)
    add_page_number(p)


def add_cover(doc, kicker, title, subtitle, footer_note):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(kicker.upper()), size=9, color=ACCENT_DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(title), font="Georgia", size=30, color=DARK, bold=False)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(subtitle), size=14, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    shade(table.cell(0, 0), DARK)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=240, start=260, bottom=240, end=260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Compliance leader. Global operator. Original thinker."), font="Georgia", size=14, color=WHITE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(84)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(footer_note), size=10, color=MUTED, italic=True)
    doc.add_page_break()


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(text.upper()), size=8.5, color=ACCENT_DARK, bold=True)


def start_page(doc, kicker):
    """Start a section on a new page without creating a trailing blank page."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.add_run().add_break(WD_BREAK.PAGE)
    set_run(p.add_run(kicker.upper()), size=8.5, color=ACCENT_DARK, bold=True)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    shade(table.cell(0, 0), "F0F3DF")
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(text), font="Georgia", size=13, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_two_col_table(doc, rows, label=""):
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(label), size=8.5, color=MUTED, italic=True)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for idx, heading in enumerate(("Prompt", "Working response / evidence")):
        shade(table.cell(0, idx), SOFT)
        p = table.cell(0, idx).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(heading), size=9.5, color=DARK, bold=True)
    for prompt, response in rows:
        cells = table.add_row().cells
        set_table_geometry(table, [2700, 6660])
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(prompt), size=9.5, color=INK, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(response), size=9.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_headshot(doc, path, caption):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(2.15))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(caption), size=8.5, color=MUTED, italic=True)


def build_media_kit():
    doc = Document()
    configure_document(doc, "Speaker & Media Kit")
    add_cover(
        doc,
        "Speaker & Media Kit · 2026",
        "Micheal Sheehy",
        "Global payments, financial crime, AI governance and inclusive leadership",
        "Prepared for event organizers, journalists, podcast hosts and communications teams",
    )

    add_kicker(doc, "At a glance")
    doc.add_heading("A practical voice on complex risk", level=1)
    doc.add_paragraph(
        "Micheal Sheehy is a global payments and compliance executive with more than 20 years of experience "
        "building financial-crime, regulatory and operating capabilities across international markets. His work "
        "connects compliance, technology, customer experience and strategy with the human realities of leadership—"
        "especially where static governance is being overtaken by faster-moving risk."
    )
    add_callout(doc, "Compliance should move as fast as the risk it manages.")
    add_two_col_table(doc, [
        ("Best for", "Keynotes, fireside chats, panels, podcasts, board and executive sessions, media commentary."),
        ("Core themes", "Adaptive compliance; global payments; AI and model governance; financial crime; inclusive leadership; LGBTQ+ visibility."),
        ("Style", "Direct, evidence-led and practical—with enough humor to keep the governance slide from becoming a medical emergency."),
        ("Location", "New York / global engagements."),
        ("Contact", "micheal@michealsheehy.com  ·  +1 (818) 397-8890  ·  michealsheehy.com"),
    ])
    doc.add_heading("Short bio", level=2)
    doc.add_paragraph(
        "Micheal Sheehy is a global payments and compliance executive with more than 20 years of experience "
        "building financial-crime, regulatory and operating capabilities across international markets. A Chief "
        "Compliance Officer and author of the Adaptive Compliance series, he writes and speaks about artificial "
        "intelligence, payments, model governance and the operating models required to scale safely. Micheal also "
        "draws on his experience as a gay executive to explore authenticity, psychological safety and why inclusive "
        "leadership improves challenge, escalation and decision-making."
    )

    start_page(doc, "Long biography")
    doc.add_heading("A global operator—and a deeply human leader", level=1)
    doc.add_paragraph(
        "Micheal Sheehy is a global payments and compliance executive with more than 20 years of experience across "
        "financial services and cross-border payments. As a Chief Compliance Officer, he has built and led large, "
        "distributed teams and major transformation programs spanning financial crime, sanctions, customer "
        "lifecycle management, transaction monitoring, fraud, artificial intelligence and model governance."
    )
    doc.add_paragraph(
        "His work sits where regulation, technology, customer experience and operating strategy meet. He is "
        "particularly interested in the gap between a control that looks impressive in a presentation and one "
        "that continues to work when the data is incomplete, the threat changes and the business is moving. His "
        "talks turn that tension into practical questions for boards, executives and operators: who owns the "
        "decision, what evidence travels with it, what should remain human and how quickly can the system learn?"
    )
    doc.add_paragraph(
        "Born in the United Kingdom, raised in Ireland and now based in New York, Micheal brings a genuinely global "
        "perspective to leadership. Growing up gay in Ireland when homosexuality was still criminalized taught him "
        "early how to read a room, how much energy people spend editing themselves and how humor can open a door "
        "that hierarchy keeps closed. Those experiences now shape his view that inclusion is not a side program. "
        "It is part of how organizations hear bad news early, make dissent safe and improve the quality of judgment."
    )
    doc.add_paragraph(
        "Micheal is the author of the eight-volume Adaptive Compliance series, which argues that compliance should "
        "operate as an adaptive system rather than a collection of periodic controls. He writes and speaks on AI-"
        "enabled financial crime, payment transparency, stablecoins, global compliance operating models, "
        "transformation and the personal realities of senior accountability. His style is direct, evidence-led and "
        "practical, with enough wit to remind an audience that governance is a serious subject—not a punishment."
    )
    add_callout(doc, "The strongest cultures make it safe to speak before certainty—and impossible to hide behind silence.")

    start_page(doc, "Keynote topics")
    doc.add_heading("Signature sessions", level=1)
    topics = [
        ("Compliance Should Move as Fast as the Risk It Manages",
         "Why periodic review and static governance are failing—and how continuous evidence, model oversight and human accountability create an adaptive system.",
         "Audience leaves with: a five-stage maturity model, practical governance questions and a clearer distinction between automation and adaptability."),
        ("While Compliance Debates AI Governance, Criminals Are Already in Production",
         "How agents, deepfakes and synthetic identities are changing the production model of financial crime—and what defensive operating models must do differently.",
         "Audience leaves with: an adversarial AI frame, authenticity controls and a roadmap for integrating fraud, cyber, KYC and financial-crime intelligence."),
        ("The Message Is the Control",
         "How payment transparency, structured data and emerging regulatory expectations are moving financial-crime decisions upstream.",
         "Audience leaves with: a payment-data quality stack, practical readiness questions and a grounded view of virtual-account complexity."),
        ("Global Compliance Without the Global Bottleneck",
         "How to combine one global standard with local expertise, follow-the-sun execution and decision-making closer to customers.",
         "Audience leaves with: decision-right principles, a capability placement model and a way to diagnose concentration risk."),
        ("Authenticity Is a Leadership Control",
         "What growing up gay in Ireland taught Micheal about reading rooms, using humor and the organizational cost of making people edit themselves.",
         "Audience leaves with: practical ways to make challenge and escalation safer, without lowering standards or accountability."),
    ]
    for title, description, outcome in topics:
        doc.add_heading(title, level=2)
        doc.add_paragraph(description)
        p = doc.add_paragraph()
        set_run(p.add_run(outcome), size=10, color=ACCENT_DARK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    set_run(p.add_run("Also available: "), size=9.5, color=DARK, bold=True)
    set_run(
        p.add_run("compliance transformation, stablecoins, de-risking, sanctions, model management and executive accountability."),
        size=9.5,
        color=MUTED,
    )

    start_page(doc, "Session design")
    doc.add_heading("Formats that work", level=1)
    add_two_col_table(doc, [
        ("20–30 minute keynote", "A sharp argument, memorable framework and practical close. Best for leadership and industry conferences."),
        ("45–60 minute keynote", "A deeper evidence-led narrative with case examples and audience reflection."),
        ("Fireside chat", "Conversational exploration of current developments, leadership decisions and lessons learned."),
        ("Executive or board session", "A private, interactive discussion centered on strategic questions, operating choices and oversight."),
        ("Panel or moderator role", "Clear synthesis, constructive provocation and enough discipline to keep six panelists from answering one question for 40 minutes."),
        ("Podcast / media interview", "Concise, quotable commentary that connects current events to durable operating implications."),
    ])
    doc.add_heading("Sample moderator introduction", level=2)
    add_callout(
        doc,
        "Micheal Sheehy is a global payments and compliance executive with more than 20 years of experience "
        "building financial-crime and regulatory capabilities across international markets. He is the author of "
        "the Adaptive Compliance series and writes about AI, payments, model governance, inclusive leadership and "
        "the choices required to scale safely. Please welcome Micheal Sheehy."
    )
    doc.add_heading("Suggested interview questions", level=2)
    add_bullets(doc, [
        "What has the compliance industry optimized that matters less than it thinks?",
        "How should boards distinguish a genuine AI control from an impressive demonstration?",
        "What changes when criminals use agents and synthetic identities at scale?",
        "What did growing up gay in Ireland teach you about silence, belonging and the duty of a leader?",
    ])

    start_page(doc, "Selected appearances")
    doc.add_heading("Media and industry conversations", level=1)
    add_two_col_table(doc, [
        ("Nasdaq TradeTalks", "How financial institutions are using AI and machine learning to combat fraud."),
        ("PYMNTS", "Payments, financial crime and the role of compliance in responsible growth."),
        ("Integrity Innovators", "Leadership, integrity and building a modern culture of compliance."),
        ("FinTech Futures", "Regulatory change, global payments and emerging financial-crime risk."),
        ("Finextra", "North Korean financial crime and the evolving role of payment platforms."),
        ("Money20/20 and industry events", "Panels and interviews on payments, compliance transformation and AI."),
    ])
    doc.add_heading("Editorial platforms", level=2)
    doc.add_paragraph(
        "Micheal’s writing at michealsheehy.com includes the eight-volume Adaptive Compliance series, essays on "
        "North Korean IT-worker networks, global KYC, stablecoins, payment transparency, de-risking, AI-enabled "
        "financial crime, inclusive leadership and the experience of finding an authentic executive voice."
    )
    add_callout(doc, "For current links, video and article references, visit michealsheehy.com/media.html.")

    start_page(doc, "Approved images")
    doc.add_heading("Headshots", level=1)
    doc.add_paragraph("High-resolution files are available on request. The images below are suitable for speaker pages, event listings and editorial use.")
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    shots = [
        (MEDIA_DIR / "micheal-sheehy-executive-closeup.jpg", "Executive close-up"),
        (MEDIA_DIR / "micheal-sheehy-formal-portrait.jpg", "Formal portrait"),
        (MEDIA_DIR / "micheal-sheehy-editorial-portrait.jpg", "Editorial portrait"),
    ]
    for idx, (path, caption) in enumerate(shots):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if path.exists():
            run.add_picture(str(path), width=Inches(1.75))
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(caption), size=8.5, color=MUTED, italic=True)
    doc.add_heading("Image and biography use", level=2)
    add_bullets(doc, [
        "Credit: Courtesy of Micheal Sheehy.",
        "Do not materially alter the images beyond normal cropping and color correction.",
        "Use the short bio for agendas and the long bio for event pages, media notes and formal introductions.",
        "Confirm current title and event-specific language before publication.",
    ])

    start_page(doc, "Contact")
    doc.add_heading("Bring the conversation to your audience", level=1)
    doc.add_paragraph(
        "For keynotes, panels, fireside chats, podcasts, board discussions, media commentary and selected advisory conversations:"
    )
    add_callout(doc, "micheal@michealsheehy.com  ·  +1 (818) 397-8890  ·  www.michealsheehy.com")
    add_bullets(doc, [
        "LinkedIn: linkedin.com/in/micheal-s-88479045",
        "Speaking: michealsheehy.com/speaking.html",
        "Media: michealsheehy.com/media.html",
        "Adaptive Compliance: michealsheehy.com/adaptive-compliance.html",
    ])
    doc.add_paragraph(
        "Please include the event or publication, audience, proposed format, timing and any deadline. "
        "All participation is subject to availability and appropriate approvals."
    )
    path = MEDIA_DIR / "Micheal_Sheehy_Speaker_Media_Kit.docx"
    doc.save(path)
    return path


def add_tool_intro(doc, number, title, purpose, use_when):
    add_kicker(doc, f"Tool {number:02d}")
    doc.add_heading(title, level=1)
    doc.add_paragraph(purpose)
    p = doc.add_paragraph()
    set_run(p.add_run("Use when: "), size=10, color=ACCENT_DARK, bold=True)
    set_run(p.add_run(use_when), size=10, color=MUTED)


def build_toolkit():
    doc = Document()
    configure_document(doc, "Adaptive Compliance Practical Toolkit")
    add_cover(
        doc,
        "Adaptive Compliance · Practical Toolkit",
        "From argument to action",
        "Seven diagnostics and checklists for modern compliance leaders",
        "Micheal Sheehy  ·  2026 edition  ·  Vendor-neutral",
    )
    add_kicker(doc, "How to use this pack")
    doc.add_heading("A working document, not a maturity certificate", level=1)
    doc.add_paragraph(
        "These tools are designed to improve the quality of leadership conversations. They do not replace legal "
        "analysis, regulatory interpretation, validation or professional judgment. Use them to expose assumptions, "
        "name ownership and identify where evidence is too weak for the decision being made."
    )
    add_callout(doc, "A high score is less valuable than an honest disagreement about why the score exists.")
    add_numbered(doc, [
        "Choose one tool and one accountable facilitator.",
        "Ask each participant to score independently before the discussion.",
        "Require evidence for every green or “mature” answer.",
        "Record the decision, owner, due date and trigger for reassessment.",
        "Repeat after a material change in product, market, model, regulation or threat.",
    ])
    add_two_col_table(doc, [
        ("1 — Unclear", "No shared standard, limited evidence or ownership not named."),
        ("2 — Emerging", "Some practice exists but depends on individuals or manual interpretation."),
        ("3 — Defined", "Documented standard and owner; application is reasonably consistent."),
        ("4 — Monitored", "Performance and exceptions are visible; thresholds trigger action."),
        ("5 — Adaptive", "The control learns from evidence and changes through governed feedback."),
    ], label="Common five-point scale used throughout the toolkit.")

    doc.add_page_break()
    add_tool_intro(doc, 1, "Payment Transparency Readiness Assessment",
                   "Test whether payment data can support enforceable transparency and earlier financial-crime decisions.",
                   "Preparing for richer structured messages, ISO 20022, virtual-account complexity or cross-border payment growth.")
    add_two_col_table(doc, [
        ("Identity integrity", "Are originator, beneficiary and relevant intermediaries represented as structured, attributable data rather than free text?  Score: ___  Evidence: __________________"),
        ("Purpose", "Can the institution explain the commercial or personal purpose of the payment and identify when the narrative conflicts with behavior?  Score: ___"),
        ("Context", "Do customer, account, corridor, counterparty and historical signals travel into the decision?  Score: ___"),
        ("Data continuity", "What material information is lost, truncated or transformed between initiation, routing and settlement?  Score: ___"),
        ("Virtual accounts", "Can the issued identifier be reliably linked to the underlying customer, owner, use case and responsible institution?  Score: ___"),
        ("Control placement", "Which risks can be addressed before execution, and which still depend on downstream monitoring?  Score: ___"),
        ("Exception governance", "Are missing, malformed and inconsistent fields measured, assigned and remediated?  Score: ___"),
        ("Regulatory traceability", "Can the organization evidence how its data and controls meet applicable transparency requirements?  Score: ___"),
    ])
    doc.add_heading("Decision", level=2)
    doc.add_paragraph("Top three weaknesses: ____________________   Named owner: ____________________   Reassessment date: ____________________")

    doc.add_page_break()
    add_tool_intro(doc, 2, "AI Governance Checklist",
                   "Put purpose, authority, evidence and accountability around AI before the system earns trust.",
                   "Approving copilots, decision-support models or agents that can call tools and affect regulated workflows.")
    add_bullets(doc, [
        "Purpose is specific enough to test, and prohibited uses are explicit.",
        "Training, retrieval and operational data have documented provenance and permitted use.",
        "The model or agent has a named business owner and independent challenge appropriate to impact.",
        "Authority is graduated: recommend, prepare, execute with approval, or execute autonomously.",
        "Human review is meaningful, not a ceremonial click after the action is effectively complete.",
        "Outputs, tool calls, prompts, versions and overrides produce reproducible evidence.",
        "Testing covers accuracy, bias, hallucination, prompt injection, data leakage and adversarial manipulation.",
        "Monitoring distinguishes model drift, workflow failure and changes in the population being assessed.",
        "Incident response includes containment, rollback, notification and learning.",
        "Third-party concentration, exit and model-change risk are understood.",
    ])
    add_two_col_table(doc, [
        ("Highest-consequence action", "____________________________________________________________"),
        ("Human approval point", "____________________________________________________________"),
        ("Evidence retained", "____________________________________________________________"),
        ("Stop / rollback trigger", "____________________________________________________________"),
        ("Accountable executive", "____________________________________________________________"),
    ])

    doc.add_page_break()
    add_tool_intro(doc, 3, "Model Drift Dashboard",
                   "Make model health visible as a set of decisions, not a collection of interesting charts.",
                   "Governing transaction monitoring, fraud, screening, customer-risk or prioritization models.")
    add_two_col_table(doc, [
        ("Data drift", "Missingness, schema changes, field distributions, latency. Signal: ______ Threshold: ______ Action: ______ Owner: ______"),
        ("Population drift", "Customer, product, geography, corridor and channel mix. Signal: ______ Threshold: ______ Action: ______"),
        ("Concept drift", "Relationship between behavior and risk outcome. Evidence: typologies, cases, investigations, external intelligence."),
        ("Performance drift", "Precision, recall proxies, coverage, stability, outcome quality and segment performance."),
        ("Operational drift", "Backlog, aging, override rates, investigator behavior, quality findings and control workarounds."),
        ("Adversarial drift", "Patterns suggesting probing, mimicry, coordinated evasion or synthetic evidence."),
    ])
    add_callout(doc, "Every metric should answer: what decision changes when this signal crosses its threshold?")

    doc.add_page_break()
    add_tool_intro(doc, 4, "Compliance Transformation Scorecard",
                   "Separate technology implementation from actual transformation.",
                   "Selecting a platform, planning a migration, assessing a troubled program or preparing for scale.")
    add_two_col_table(doc, [
        ("Data readiness", "Quality, lineage, access, ownership and suitability for intended decisions.  Score: ___"),
        ("Process redesign", "Legacy steps removed; handoffs, exceptions and controls designed end to end.  Score: ___"),
        ("Model governance", "Inventory, validation, monitoring, change control and evidence.  Score: ___"),
        ("Adoption", "Role-based training, workflow fit, behavior measures and feedback.  Score: ___"),
        ("Operating ownership", "Named permanent owner, decision rights and service model.  Score: ___"),
        ("Coexistence", "Old/new reconciliation, historical evidence, cutover and rollback controls.  Score: ___"),
        ("Outcome measures", "Risk coverage, customer friction, time to adapt and decision quality.  Score: ___"),
    ])
    doc.add_paragraph("Lowest-scoring dimension: ____________________   Decision required: ____________________   Owner: ____________________")

    doc.add_page_break()
    add_tool_intro(doc, 5, "Global Operating Model Diagnostic",
                   "Place standards, capabilities and judgment where they create the strongest combination of consistency and speed.",
                   "Designing global hubs, follow-the-sun coverage or correcting excessive centralization.")
    add_two_col_table(doc, [
        ("Standards", "What must be globally consistent? Who owns the policy, data definitions and minimum outcome?"),
        ("Local judgment", "Which decisions require language, market, customer or regulatory proximity?"),
        ("Shared capabilities", "Which specialist skills should be global centers of excellence rather than duplicated everywhere?"),
        ("Decision rights", "Where can regions decide, where must they consult and what must escalate?"),
        ("Handoffs", "What evidence must travel with work across time zones, and who owns it during transition?"),
        ("Concentration", "Which people, locations, vendors or datasets create single points of failure?"),
        ("Calibration", "How are differences detected, debated and converted into shared learning?"),
        ("Leadership depth", "Can each hub lead and challenge, or is it only an execution location?"),
    ])
    add_callout(doc, "Centralize the standard. Distribute informed judgment. Connect the system through evidence.")

    doc.add_page_break()
    add_tool_intro(doc, 6, "Board Questions for AI Agents",
                   "Help directors challenge agentic systems at the level of consequence and accountability.",
                   "Board approval, risk committee oversight or a material expansion of agent authority.")
    add_numbered(doc, [
        "What can the agent do—not merely recommend—and what is the highest-consequence action available to it?",
        "Which decisions are reversible, and how quickly can the organization stop or roll back the system?",
        "What evidence allows an independent reviewer to reproduce the agent’s reasoning, tool calls and actions?",
        "Where does human approval occur, and is the reviewer given enough time and information to disagree?",
        "How is the agent tested against manipulation, unexpected tool behavior, compromised data and adversarial users?",
        "What changes would require reapproval: model, data source, tool, purpose, authority or market?",
        "How will the board distinguish productivity gains from risk merely moved elsewhere?",
        "Who is personally accountable when the agent follows its instructions and still causes harm?",
        "What third parties or foundation models could create concentration, continuity or exit risk?",
        "What incident would cause management to notify the board immediately?",
    ])
    doc.add_paragraph("Board action / challenge: __________________________________________________________________________")

    doc.add_page_break()
    add_tool_intro(doc, 7, "30/60/90-Day Transformation Plan",
                   "Create momentum without confusing speed with ungoverned activity.",
                   "Starting or resetting a compliance transformation.")
    doc.add_heading("Days 1–30 · Establish the truth", level=2)
    add_bullets(doc, [
        "Name the accountable executive, business owner and decision forum.",
        "Baseline risk coverage, data quality, process performance, customer friction and operating pain.",
        "Map material dependencies, legacy systems, manual work and current exceptions.",
        "Choose one outcome that can be improved and evidenced within 90 days.",
    ])
    doc.add_heading("Days 31–60 · Prove the change", level=2)
    add_bullets(doc, [
        "Redesign the target process before automating it.",
        "Test data, model and workflow assumptions with real users and representative cases.",
        "Define coexistence, rollback, quality and incident controls.",
        "Train the people who will own the capability—not only the implementation team.",
    ])
    doc.add_heading("Days 61–90 · Govern the scale", level=2)
    add_bullets(doc, [
        "Compare outcome evidence to the baseline and challenge unintended consequences.",
        "Move ownership into the permanent operating model.",
        "Approve the next scale decision with explicit thresholds and stop conditions.",
        "Publish the learning: what changed, what did not and what the next review must test.",
    ])
    add_two_col_table(doc, [
        ("90-day outcome", "____________________________________________________________"),
        ("Executive owner", "____________________________________________________________"),
        ("Evidence of improvement", "____________________________________________________________"),
        ("Scale / stop decision date", "____________________________________________________________"),
    ])

    doc.add_page_break()
    add_kicker(doc, "Close")
    doc.add_heading("The tool is the conversation it makes possible", level=1)
    doc.add_paragraph(
        "Adaptive compliance is not a product or a finish line. It is the capability to see change, interpret "
        "evidence, make accountable decisions and improve the control environment while risk is still moving."
    )
    add_callout(doc, "Better compliance begins when the organization can change its mind for the right reason—and show its work.")
    doc.add_paragraph("More writing, tools and the eight-volume series: www.michealsheehy.com")
    doc.add_paragraph("Contact: micheal@michealsheehy.com  ·  +1 (818) 397-8890")
    path = TOOLS_DIR / "Adaptive_Compliance_Practical_Toolkit.docx"
    doc.save(path)
    return path


def build_toolkit():
    """Build the detailed executive toolkit maintained in its dedicated generator."""
    try:
        from scripts.generate_adaptive_toolkit import build_toolkit as build_detailed_toolkit
    except ImportError:
        from generate_adaptive_toolkit import build_toolkit as build_detailed_toolkit
    return build_detailed_toolkit()


if __name__ == "__main__":
    media = build_media_kit()
    toolkit = build_toolkit()
    print(media)
    print(toolkit)
