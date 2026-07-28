from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "assets" / "tools"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Compact Reference Guide preset, with a named Micheal Sheehy brand override.
DARK = "192720"
INK = "17211D"
MUTED = "66716B"
ACCENT = "C5DB49"
ACCENT_DARK = "738125"
SOFT = "E5EBE4"
PALE = "F4F1EA"
WHITE = "FFFFFF"
LINE = "D7DDD8"
GOLD = "9A7416"
GOLD_FILL = "F6EED8"
RED = "9B3B31"
RED_FILL = "F7E7E4"
BLUE = "315C6D"
BLUE_FILL = "E5EEF1"


def set_run(run, *, font="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_cell_margins(cell, top=55, start=110, bottom=55, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, end])


def configure_document(doc):
    section = doc.sections[0]
    # Named density override: the workbook uses a tighter vertical canvas than
    # the base compact-reference preset while retaining the 6.5-inch text width.
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.34)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, before, after, color in (
        ("Heading 1", 15, 12, 6, DARK),
        ("Heading 2", 12, 9, 4, DARK),
        ("Heading 3", 10.5, 7, 3, ACCENT_DARK),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.1

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("ADAPTIVE COMPLIANCE · EXECUTIVE TOOLKIT"), size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    set_run(p.add_run("MICHEAL SHEEHY  ·  "), size=8.5, color=MUTED, bold=True)
    add_page_number(p)


def start_page(doc, kicker, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(kicker.upper()), size=8.5, color=ACCENT_DARK, bold=True)
    doc.add_heading(title, level=1)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        set_run(p.add_run(subtitle), size=10, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("ADAPTIVE COMPLIANCE · EXECUTIVE TOOLKIT"), size=9, color=ACCENT_DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("From argument to operating discipline"), font="Georgia", size=28, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        p.add_run("Seven executive diagnostics for payments, AI, models, transformation and global compliance"),
        size=13,
        color=MUTED,
    )

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    labels = [
        ("DIAGNOSE", "Expose where evidence is weak"),
        ("DECIDE", "Name the accountable choice"),
        ("ADAPT", "Set the trigger to reassess"),
    ]
    for idx, (label, detail) in enumerate(labels):
        cell = table.cell(0, idx)
        shade(cell, DARK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(label), size=9, color=ACCENT, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(detail), size=9.5, color=WHITE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Micheal Sheehy"), font="Georgia", size=16, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("2026 edition · Vendor-neutral · Designed for executive working sessions"), size=9.5, color=MUTED, italic=True)


def add_callout(doc, label, text, *, fill=PALE, label_color=ACCENT_DARK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=100, start=190, bottom=100, end=190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(label.upper()), size=8.5, color=label_color, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(text), font="Georgia", size=12.5, color=DARK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        p.add_run(item)


def add_table(doc, headers, rows, widths, *, font_size=9.2, header_fill=SOFT, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for idx, heading in enumerate(headers):
        cell = table.cell(0, idx)
        shade(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(heading), size=font_size, color=WHITE if header_fill == DARK else DARK, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        set_table_geometry(table, widths)
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.02
            set_run(
                p.add_run(value),
                size=font_size,
                color=INK if idx == 0 else MUTED,
                bold=first_col_bold and idx == 0,
            )
    return table


def add_tool_header(doc, number, title, decision, use_when, output):
    start_page(doc, f"Tool {number:02d}", title, decision)
    add_table(
        doc,
        ["USE WHEN", "PRIMARY OUTPUT"],
        [(use_when, output)],
        [4680, 4680],
        font_size=9.5,
        header_fill=DARK,
    )


def add_decision_record(doc, fields):
    doc.add_heading("Executive decision record", level=2)
    rows = [(label, "________________________________________________________________________________") for label in fields]
    add_table(doc, ["DECISION FIELD", "WORKING RECORD"], rows, [2400, 6960], font_size=9.2, first_col_bold=True)


def add_action_register(doc, rows=4):
    doc.add_heading("Action register", level=2)
    data = [(" \n ", " \n ", " \n ", " \n ", " \n ") for _ in range(rows)]
    add_table(
        doc,
        ["ACTION", "OWNER", "DUE", "SUCCESS MEASURE", "REASSESSMENT TRIGGER"],
        data,
        [2600, 1350, 900, 2200, 2310],
        font_size=8.3,
        header_fill=DARK,
    )


def build_toolkit():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    start_page(
        doc,
        "Operating guide",
        "How to use the toolkit",
        "Use each tool to convert a broad risk discussion into an evidence-backed decision with a named owner.",
    )
    add_callout(
        doc,
        "The standard",
        "A score is not the output. The output is a decision, the evidence behind it and the condition that would make the organization reconsider.",
    )
    doc.add_heading("Choose the right starting point", level=2)
    add_table(
        doc,
        ["TRIGGER", "START WITH", "DECISION SUPPORTED"],
        [
            ("Richer payment messages, ISO 20022 or virtual accounts", "01 · Payment Transparency", "Where data breaks undermine identity, purpose or traceability."),
            ("Copilot, model or agent entering a regulated workflow", "02 · AI & Agent Governance", "What authority is permitted and what evidence is required before release."),
            ("Model performance is changing or outcomes are disputed", "03 · Model Health & Drift", "Whether to continue, tune, constrain, validate or stop the model."),
            ("New platform, migration or program reset", "04 · Transformation Readiness", "Whether the program is ready to move through the next gate."),
            ("Global hubs, regional growth or central bottlenecks", "05 · Global Operating Model", "Which capabilities and decisions belong globally, regionally or locally."),
            ("Board oversight of high-impact AI", "06 · Board Oversight", "What the board should approve, challenge or require management to evidence."),
            ("Transformation needs traction in one quarter", "07 · 30/60/90-Day Plan", "What will change, who owns it and what proves improvement."),
        ],
        [3150, 2200, 4010],
        font_size=8.7,
        first_col_bold=True,
    )
    doc.add_heading("Working-session discipline", level=2)
    add_bullets(doc, [
        "Pre-read: circulate the relevant tool, baseline data and known issues at least two working days in advance.",
        "Score independently: collect individual views before the meeting so hierarchy does not manufacture agreement.",
        "Challenge the evidence: require a source, owner and date for every strong claim or high score.",
        "Close explicitly: record the decision, dissent, residual risk, owner, due date and reassessment trigger.",
    ])

    start_page(
        doc,
        "Operating guide",
        "Scoring, confidence and evidence",
        "A common standard prevents teams from using the same score to describe very different realities.",
    )
    add_table(
        doc,
        ["LEVEL", "ANCHOR", "WHAT EXISTS", "EVIDENCE EXPECTATION", "LEADERSHIP IMPLICATION"],
        [
            ("1", "Unclear", "No shared standard; ownership or intended outcome is disputed.", "Anecdote, incomplete inventory or no reliable record.", "Do not scale. Establish ownership and the minimum standard."),
            ("2", "Emerging", "Practice exists but depends on individuals or local workarounds.", "Examples exist; coverage and consistency are unknown.", "Stabilize the process and expose variation."),
            ("3", "Defined", "Documented standard, owner and workflow; application is reasonably consistent.", "Policy, process, inventory and sampled execution evidence.", "Prove performance and manage exceptions."),
            ("4", "Monitored", "Performance, exceptions and thresholds are visible; action is governed.", "Current metrics, trends, thresholds, decisions and closure evidence.", "Test whether signals drive timely decisions."),
            ("5", "Adaptive", "The control learns from evidence and changes through accountable feedback.", "Closed-loop evidence showing detection, decision, change and validation.", "Protect against complacency and uncontrolled optimization."),
        ],
        [850, 1200, 2250, 2550, 2510],
        font_size=8.1,
        first_col_bold=True,
    )
    doc.add_heading("Confidence is separate from maturity", level=2)
    add_table(
        doc,
        ["CONFIDENCE", "USE WHEN", "REQUIRED RESPONSE"],
        [
            ("High", "Evidence is current, representative, attributable and independently challengeable.", "Use the score; retain the evidence reference."),
            ("Medium", "Evidence is directionally useful but incomplete, stale or limited to selected populations.", "Use provisionally; create a dated evidence gap action."),
            ("Low", "The score depends on assertion, a narrow sample or an owner’s judgment.", "Treat as unknown; do not use the score to justify expansion."),
        ],
        [1200, 4500, 3660],
        font_size=9,
        header_fill=BLUE_FILL,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Facilitator challenge",
        "Where would an independent reviewer look to reproduce this answer—and what might they conclude differently?",
        fill=GOLD_FILL,
        label_color=GOLD,
    )
    add_decision_record(doc, ["Scope assessed", "Facilitator", "Participants", "Decision forum", "Evidence cut-off date"])

    # TOOL 01
    add_tool_header(
        doc,
        1,
        "Payment Transparency Readiness",
        "Determine whether payment data can support identity, purpose and contextual decisions before risk disappears into the transaction chain.",
        "Preparing for structured payment data, cross-border growth, virtual-account issuance or stronger transparency enforcement.",
        "A data-break heatmap, control-placement decision and accountable remediation plan.",
    )
    doc.add_heading("Follow the data, not the organizational chart", level=2)
    add_table(
        doc,
        ["INITIATION", "VALIDATION", "ROUTING", "SCREENING", "SETTLEMENT", "MONITORING"],
        [("Who and why enter the message.", "Required fields and format are checked.", "Data is transformed and passed.", "Risk is assessed before release.", "Value and message reach destination.", "Behavior and outcomes feed learning.")],
        [1560, 1560, 1560, 1560, 1560, 1560],
        font_size=7.9,
        header_fill=DARK,
    )
    doc.add_heading("Five control outcomes", level=2)
    add_table(
        doc,
        ["OUTCOME", "EXECUTIVE TEST", "FAILURE SIGNAL"],
        [
            ("Attributable identity", "Can the institution identify the parties, intermediaries and responsible customer behind the payment?", "Alias, free text, pooled identity or an issued account identifier cannot be resolved."),
            ("Meaningful purpose", "Can the stated purpose be compared with customer profile, product and behavior?", "Narrative is generic, inconsistent or unavailable to the decision point."),
            ("Preserved context", "Does relevant information survive transformation, routing and settlement?", "Fields are truncated, overwritten, mapped inconsistently or dropped across rails."),
            ("Timely intervention", "Which risks can be addressed before execution rather than discovered downstream?", "The decisive information exists but reaches monitoring only after value moves."),
            ("Traceable accountability", "Can every exception be assigned to an institution, control owner and remediation path?", "Missing data is tolerated without measurement, consequence or feedback."),
        ],
        [1700, 4200, 3460],
        font_size=8.7,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Virtual-account complexity",
        "A virtual account is not merely another account number. The control must preserve the link between the issued identifier, underlying customer, beneficial ownership, permitted purpose and institution responsible for monitoring.",
        fill=BLUE_FILL,
        label_color=BLUE,
    )
    doc.add_heading("Immediate red flags", level=2)
    add_bullets(doc, [
        "The same payment can be attributed differently by onboarding, operations, sanctions and transaction monitoring.",
        "Message-quality exceptions are measured as processing defects but not evaluated as control failures.",
        "Virtual-account ownership can be reconstructed only through manual investigation or local knowledge.",
        "Commercial pressure can override missing-data rules without a documented risk acceptance.",
    ])

    start_page(
        doc,
        "Tool 01 · Diagnostic",
        "Payment transparency maturity matrix",
        "Score each dimension 1–5, record confidence and cite the evidence. The anchors below distinguish weak, defined and adaptive practice.",
    )
    add_table(
        doc,
        ["DIMENSION", "1 · UNCLEAR", "3 · DEFINED", "5 · ADAPTIVE"],
        [
            ("Identity integrity", "Party roles and identifiers vary by rail or system.", "Common definitions and validation exist for material flows.", "Identity quality is monitored end to end; breaks change controls and upstream design."),
            ("Purpose", "Purpose is optional, generic or free text only.", "Purpose standards exist for priority products and corridors.", "Purpose is structured, tested against context and improved through investigation outcomes."),
            ("Data structure", "Material fields rely on free text or local mapping.", "Required fields, formats and ownership are documented.", "Structured data is governed as a control asset with quality thresholds and lineage."),
            ("Continuity", "Loss or transformation across the chain is unknown.", "Material transformations and handoffs are mapped.", "Automated reconciliation detects loss, mutation and unexpected routing changes."),
            ("Virtual accounts", "Issued identifiers are not consistently linked to the underlying customer.", "Ownership, purpose and issuer responsibility are documented.", "Linkage is real-time, testable and available to screening, monitoring and investigations."),
            ("Control placement", "Controls run where systems happen to permit them.", "Pre- and post-execution controls are intentionally assigned.", "Controls move upstream when evidence shows earlier action improves outcomes."),
            ("Exception governance", "Missing data is handled case by case.", "Thresholds, owners and remediation paths exist.", "Exception trends trigger product, counterparty or rail-level decisions."),
            ("Regulatory traceability", "Requirements are interpreted in disconnected projects.", "Applicable obligations map to data and controls.", "Changes in expectations trigger governed reassessment and evidence refresh."),
        ],
        [1680, 2560, 2560, 2560],
        font_size=7.75,
        first_col_bold=True,
    )
    doc.add_heading("Scorecard", level=2)
    add_table(
        doc,
        ["DIMENSION", "SCORE 1–5", "CONFIDENCE H/M/L", "EVIDENCE REFERENCE", "OWNER"],
        [(name, "", "", "", "") for name in [
            "Identity integrity", "Purpose", "Data structure", "Continuity",
            "Virtual accounts", "Control placement", "Exception governance", "Regulatory traceability"
        ]],
        [1800, 1050, 1550, 3260, 1700],
        font_size=8.3,
        header_fill=DARK,
        first_col_bold=True,
    )

    start_page(
        doc,
        "Tool 01 · Decision",
        "From data breaks to accountable action",
        "Prioritize breaks by control consequence, not by implementation convenience.",
    )
    add_table(
        doc,
        ["BREAK / GAP", "WHERE INTRODUCED", "CONTROL IMPACT", "POPULATION / VALUE", "CURRENT COMPENSATING CONTROL", "EVIDENCE OWNER"],
        [("", "", "", "", "", "") for _ in range(6)],
        [1650, 1350, 1600, 1350, 2150, 1260],
        font_size=7.8,
        header_fill=SOFT,
    )
    doc.add_heading("Prioritization test", level=2)
    add_table(
        doc,
        ["QUESTION", "RESPONSE"],
        [
            ("What risk decision is weakened by this break?", ""),
            ("Can the information be captured or validated earlier?", ""),
            ("Who can change the source, mapping, rail or commercial rule?", ""),
            ("What compensating control exists, and how is its effectiveness proven?", ""),
            ("What threshold requires product restriction, counterparty escalation or risk acceptance?", ""),
        ],
        [3900, 5460],
        font_size=9,
        first_col_bold=True,
    )
    add_decision_record(doc, [
        "Decision and target state", "Residual risk accepted", "Accountable executive / control owner",
        "Review date", "Trigger for immediate reassessment"
    ])
    start_page(
        doc,
        "Tool 01 · Action",
        "Payment transparency delivery plan",
        "Translate the control decision into a sequenced remediation plan with measurable closure evidence.",
    )
    add_action_register(doc, rows=6)
    doc.add_heading("Progress review", level=2)
    add_table(
        doc,
        ["REVIEW FIELD", "WORKING RECORD"],
        [
            ("Review cadence and forum", ""),
            ("Evidence required for closure", ""),
            ("Dependencies requiring executive intervention", ""),
            ("Decision if delivery misses tolerance", ""),
        ],
        [3300, 6060],
        font_size=9,
        first_col_bold=True,
    )

    # TOOL 02
    add_tool_header(
        doc,
        2,
        "AI & Agent Governance",
        "Set authority, evidence and accountability before an AI system can influence a regulated outcome.",
        "Approving a copilot, decision-support model or agent that can access data, call tools or affect customers.",
        "An authority classification, control evidence pack, challenge-test record and release decision.",
    )
    doc.add_heading("Start with consequence, not sophistication", level=2)
    add_table(
        doc,
        ["AUTHORITY", "SYSTEM MAY", "EXAMPLE", "DEFAULT GOVERNANCE"],
        [
            ("0 · Prohibited", "Not operate for the proposed purpose.", "Unapproved external action or sensitive-data use.", "Block technically; record prohibited use."),
            ("1 · Inform", "Retrieve, summarize or explain.", "Policy assistant or case summary.", "Data controls, accuracy testing and user disclosure."),
            ("2 · Recommend", "Suggest a decision without preparing execution.", "Risk prioritization or next-best action.", "Outcome testing, explainability and human challenge."),
            ("3 · Prepare", "Draft the action and assemble evidence.", "Prepare an RFI, alert disposition or case note.", "Approval workflow, traceability and override monitoring."),
            ("4 · Execute with approval", "Act only after meaningful human authorization.", "Submit a customer restriction after review.", "Consequence-based approval, rollback and incident response."),
            ("5 · Bounded autonomy", "Execute within explicit constraints and limits.", "Low-risk routing or standardized remediation.", "Independent validation, continuous monitoring, kill switch and board visibility where material."),
        ],
        [1500, 2700, 2500, 2660],
        font_size=8.15,
        first_col_bold=True,
    )
    doc.add_heading("Use-case profile", level=2)
    add_table(
        doc,
        ["FIELD", "WORKING DEFINITION"],
        [
            ("Purpose and intended user", ""),
            ("Affected customers, employees or counterparties", ""),
            ("Highest-consequence available action", ""),
            ("Data accessed, retained or transmitted", ""),
            ("Tools, systems and third parties connected", ""),
            ("Reversibility and maximum time to stop", ""),
            ("Proposed authority level", ""),
        ],
        [3000, 6360],
        font_size=9,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Executive question",
        "If the system follows its instructions perfectly and still causes harm, who owns the design failure?",
        fill=GOLD_FILL,
        label_color=GOLD,
    )

    start_page(
        doc,
        "Tool 02 · Control design",
        "Lifecycle control and evidence matrix",
        "Approval should rely on an evidence pack that can survive independent challenge and future reconstruction.",
    )
    add_table(
        doc,
        ["LIFECYCLE", "CONTROL EXPECTATION", "MINIMUM EVIDENCE", "OWNER"],
        [
            ("Purpose", "Intended outcome, prohibited uses, affected population and success measure are explicit.", "Approved use-case statement; consequence assessment.", ""),
            ("Data", "Provenance, permission, quality, retention and sensitive-data handling are governed.", "Data inventory; lineage; access record; quality results.", ""),
            ("Model / prompt", "Version, configuration, retrieval sources and system instructions are controlled.", "Version record; configuration baseline; change history.", ""),
            ("Tools and authority", "Available actions are bounded to the approved level and environment.", "Tool inventory; permissions; transaction limits; kill switch.", ""),
            ("Testing", "Representative, edge, bias and adversarial cases test both output and action.", "Test set; acceptance thresholds; failure analysis; independent challenge.", ""),
            ("Human oversight", "Review occurs before consequence and gives the reviewer time and evidence to disagree.", "Workflow design; reviewer instructions; override and disagreement data.", ""),
            ("Monitoring", "Quality, behavior, population, tool calls, overrides and unintended outcomes are visible.", "Dashboard; thresholds; trend analysis; escalation record.", ""),
            ("Incident and change", "Containment, rollback, notification and reapproval triggers are predefined.", "Incident playbook; change taxonomy; rehearsal evidence.", ""),
            ("Third parties", "Concentration, model changes, continuity, exit and subcontractor risk are understood.", "Contract rights; dependency map; contingency and exit plan.", ""),
        ],
        [1500, 3500, 3200, 1160],
        font_size=7.9,
        first_col_bold=True,
    )
    doc.add_heading("Release gates", level=2)
    add_table(
        doc,
        ["GATE", "PASS CONDITION", "RESULT / EVIDENCE"],
        [
            ("Design", "Purpose, authority and affected population approved.", ""),
            ("Control", "Permissions, approval points, evidence retention and rollback implemented.", ""),
            ("Challenge", "Acceptance and adversarial tests passed or residuals explicitly accepted.", ""),
            ("Operational", "Owners, monitoring, incident response and support model are live.", ""),
            ("Change", "Events requiring revalidation or reapproval are configured and owned.", ""),
        ],
        [1500, 4400, 3460],
        font_size=8.8,
        header_fill=BLUE_FILL,
        first_col_bold=True,
    )

    start_page(
        doc,
        "Tool 02 · Challenge",
        "Adversarial test record and release decision",
        "Test what an intelligent adversary, compromised dependency or rushed operator can make the system do.",
    )
    add_table(
        doc,
        ["SCENARIO", "EXPECTED CONTROL", "RESULT", "SEVERITY", "REMEDIATION / OWNER"],
        [
            ("Prompt or instruction injection", "System rejects or isolates untrusted instructions.", "", "", ""),
            ("Compromised or manipulated retrieval data", "Source controls and anomaly detection prevent unsafe reliance.", "", "", ""),
            ("Unexpected tool behavior or unavailable dependency", "System fails safely and preserves evidence.", "", "", ""),
            ("Confidently incorrect output", "Acceptance rules or human review prevent consequence.", "", "", ""),
            ("Sensitive-data leakage", "Data boundaries, redaction and output controls contain exposure.", "", "", ""),
            ("Biased or uneven segment outcomes", "Segment testing and thresholds trigger remediation.", "", "", ""),
            ("Human rubber-stamping", "Workflow, sampling and override metrics expose ineffective review.", "", "", ""),
            ("Unauthorized authority expansion", "Permissions and change controls block new actions.", "", "", ""),
        ],
        [2300, 3000, 1050, 900, 2110],
        font_size=7.7,
        header_fill=SOFT,
        first_col_bold=True,
    )
    add_decision_record(doc, [
        "Release decision and approved authority", "Conditions before release", "Residual risk owner",
        "Monitoring / rollback threshold", "Reapproval trigger"
    ])
    add_action_register(doc, rows=3)

    # TOOL 03
    add_tool_header(
        doc,
        3,
        "Model Health & Drift",
        "Turn model monitoring into a governed decision about whether to continue, tune, constrain, validate or stop.",
        "Governing transaction monitoring, fraud, screening, customer-risk, prioritization or decision-support models.",
        "A decision-linked dashboard, drift triage record and governed response.",
    )
    add_callout(
        doc,
        "Operating principle",
        "A metric without a threshold is observation. A threshold without a decision is decoration.",
    )
    add_table(
        doc,
        ["DRIFT LENS", "WHAT CHANGES", "PRACTICAL SIGNALS", "DECISION IT MAY TRIGGER"],
        [
            ("Data", "Inputs, availability, schema, latency or missingness.", "Null rates; field distributions; data freshness; mapping errors.", "Constrain population, repair source, reject data or pause."),
            ("Population", "Customer, product, corridor, geography or channel mix.", "Mix shift; new segment share; volume/value concentration.", "Segment controls, recalibration or expanded testing."),
            ("Concept", "Relationship between behavior and the risk outcome.", "New typologies; investigation outcomes; external intelligence.", "Feature or rule change; scenario redesign; risk appetite review."),
            ("Performance", "Coverage, stability, precision or outcome quality.", "Hit quality; known-event recall; stability; segment performance.", "Tune, validate, supplement or replace."),
            ("Operational", "How people and workflows use the model.", "Backlog; aging; overrides; QA findings; workarounds.", "Capacity, workflow or threshold change."),
            ("Adversarial", "Evidence of probing, mimicry or coordinated evasion.", "Boundary clustering; repeated near misses; synthetic patterns.", "Threat response, control randomization or urgent redesign."),
        ],
        [1450, 2450, 3000, 2460],
        font_size=8.1,
        first_col_bold=True,
    )
    doc.add_heading("Monitoring design test", level=2)
    add_table(
        doc,
        ["FOR EVERY SIGNAL, NAME", "WORKING RESPONSE"],
        [
            ("Population and segment covered", ""),
            ("Data source and calculation owner", ""),
            ("Expected range and threshold", ""),
            ("Decision owner and maximum response time", ""),
            ("Action at amber and red", ""),
            ("Evidence required to close the issue", ""),
        ],
        [3600, 5760],
        font_size=9,
        first_col_bold=True,
    )

    start_page(
        doc,
        "Tool 03 · Dashboard",
        "Decision-linked model health dashboard",
        "The executive view should explain what changed, why it matters and what management is doing—not merely show more charts.",
    )
    add_table(
        doc,
        ["SIGNAL", "SEGMENT", "CURRENT", "THRESHOLD", "TREND", "STATUS", "DECISION / ACTION", "OWNER"],
        [("", "", "", "", "", "", "", "") for _ in range(10)],
        [1350, 1050, 900, 900, 700, 750, 2600, 1110],
        font_size=7.6,
        header_fill=DARK,
    )
    doc.add_heading("Minimum metric set", level=2)
    add_table(
        doc,
        ["DOMAIN", "LEADING INDICATORS", "OUTCOME / ASSURANCE INDICATORS"],
        [
            ("Input integrity", "Missingness, freshness, schema change, source reconciliation.", "Material data incidents; control populations affected."),
            ("Population", "Mix shifts by product, geography, channel, risk and customer tenure.", "Segment-level outcome divergence."),
            ("Model behavior", "Score distribution, feature availability, stability, boundary clustering.", "Known-event coverage; precision proxies; validation findings."),
            ("Operations", "Alert/case volume, aging, overrides, queue concentration, investigator variance.", "QA defects; SLA breaches; material missed or delayed action."),
            ("Threat response", "New typologies, evasion indicators, intelligence-to-control lead time.", "Detection uplift; time from signal to validated change."),
        ],
        [1700, 3820, 3840],
        font_size=8.4,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Executive challenge",
        "Which metric would make management stop the model today—and is that threshold actually instrumented?",
        fill=RED_FILL,
        label_color=RED,
    )

    start_page(
        doc,
        "Tool 03 · Response",
        "Drift triage and governed change",
        "Severity should reflect control consequence, affected population, duration and reversibility.",
    )
    add_table(
        doc,
        ["SEVERITY", "CONDITION", "INITIAL RESPONSE", "DECISION FORUM"],
        [
            ("S1 · Observe", "Within tolerance; no material outcome concern.", "Record and continue monitoring.", "Model owner."),
            ("S2 · Investigate", "Threshold breach or unexplained segment change; limited consequence.", "Validate data, isolate segment and set dated analysis.", "Model owner + compliance operations."),
            ("S3 · Constrain", "Material degradation, control gap or prolonged uncertainty.", "Apply compensating control, restrict population and notify governance.", "Model governance / risk committee."),
            ("S4 · Stop / incident", "Potential regulatory, customer or financial harm; evidence unreliable.", "Stop or rollback, preserve evidence and activate incident process.", "Accountable executive + incident governance."),
        ],
        [1400, 3300, 3100, 1560],
        font_size=8.4,
        first_col_bold=True,
    )
    doc.add_heading("Triage record", level=2)
    add_table(
        doc,
        ["FIELD", "WORKING RECORD"],
        [
            ("Signal and detection time", ""),
            ("Affected model, version, population and period", ""),
            ("Known / suspected control consequence", ""),
            ("Severity and rationale", ""),
            ("Immediate containment or compensating control", ""),
            ("Evidence required to decide", ""),
            ("Validation / reapproval requirement", ""),
            ("Closure criteria", ""),
        ],
        [3300, 6060],
        font_size=9,
        first_col_bold=True,
    )
    add_decision_record(doc, ["Decision", "Decision owner", "Dissent / challenge", "Effective date", "Review date", "Reassessment trigger"])

    # TOOL 04
    add_tool_header(
        doc,
        4,
        "Transformation Readiness",
        "Separate technology delivery from a demonstrable change in risk outcomes, customer experience and operating ownership.",
        "Selecting a platform, planning a migration, assessing a troubled program or deciding whether to scale.",
        "A readiness heatmap, stage-gate decision, coexistence plan and benefits record.",
    )
    add_table(
        doc,
        ["DIMENSION", "EXECUTIVE QUESTION", "EVIDENCE THAT SHOULD EXIST", "COMMON FAILURE"],
        [
            ("Problem definition", "What outcome must change, for whom and by how much?", "Baseline, target, scope and explicit non-goals.", "Technology becomes the strategy."),
            ("Data readiness", "Can the required decision be made from available data?", "Quality, lineage, ownership, access and representative test data.", "Poor data is discovered during migration."),
            ("Process redesign", "Which legacy steps, handoffs and controls will disappear?", "End-to-end target process and exception design.", "Old process is recreated in a new interface."),
            ("Control / model governance", "How will logic, models and changes be independently challenged?", "Inventory, validation, testing, change and monitoring plan.", "Governance begins after launch."),
            ("Adoption", "How will behavior change be observed, not merely trained?", "Role design, workflow testing, usage, override and quality measures.", "Training completion is treated as adoption."),
            ("Operating ownership", "Who owns outcomes after the project team leaves?", "Named service owner, decision rights, capacity and support model.", "Permanent ownership is deferred."),
            ("Coexistence and cutover", "How will old and new evidence reconcile, and how can the firm roll back?", "Dual-run, reconciliation, cutover thresholds and rollback plan.", "Two systems create two versions of truth."),
            ("Benefits and unintended effects", "What proves better risk and customer outcomes?", "Outcome measures, control coverage, friction and cost baselines.", "Delivery milestones substitute for benefits."),
        ],
        [1600, 3100, 2900, 1760],
        font_size=7.8,
        first_col_bold=True,
    )
    doc.add_heading("Readiness scorecard", level=2)
    add_table(
        doc,
        ["DIMENSION", "SCORE 1–5", "CONFIDENCE", "EVIDENCE GAP", "DECISION OWNER"],
        [(name, "", "", "", "") for name in [
            "Problem definition", "Data readiness", "Process redesign", "Control/model governance",
            "Adoption", "Operating ownership", "Coexistence/cutover", "Benefits"
        ]],
        [2050, 1050, 1150, 3300, 1810],
        font_size=8.2,
        header_fill=DARK,
        first_col_bold=True,
    )

    start_page(
        doc,
        "Tool 04 · Stage gates",
        "Transformation stage-gate standard",
        "A gate is a decision backed by exit evidence—not a meeting scheduled at the end of a project phase.",
    )
    add_table(
        doc,
        ["GATE", "QUESTION", "MINIMUM EXIT EVIDENCE", "DECISION"],
        [
            ("0 · Frame", "Is the problem worth solving and defined in outcome terms?", "Baseline; affected population; target; sponsor; risk of inaction.", "Fund discovery / stop."),
            ("1 · Design", "Will the target process and control model solve the problem?", "Process, data, control, operating and architecture designs; key risks.", "Approve build / redesign."),
            ("2 · Build", "Is the capability complete enough for representative testing?", "Configured solution; lineage; controls; test environment; ownership.", "Enter proving / remediate."),
            ("3 · Prove", "Does it work for risk, operations and customers under realistic conditions?", "Acceptance, control, adversarial, volume, segment and user evidence.", "Approve limited release / fail."),
            ("4 · Cut over", "Can the organization move safely and retain evidence?", "Reconciliation; coexistence; rollback; support; incident readiness.", "Cut over / extend dual run."),
            ("5 · Scale", "Are outcomes sustained and unintended effects controlled?", "Outcome trend; defects; adoption; capacity; validation; residual risk.", "Scale / constrain / stop."),
        ],
        [1300, 2750, 3750, 1560],
        font_size=8.1,
        first_col_bold=True,
    )
    doc.add_heading("Gate decision record", level=2)
    add_table(
        doc,
        ["FIELD", "WORKING RECORD"],
        [
            ("Gate and date", ""),
            ("Evidence reviewed", ""),
            ("Acceptance criteria met / missed", ""),
            ("Material dissent", ""),
            ("Decision: proceed / conditional / hold / stop", ""),
            ("Conditions and accountable owners", ""),
            ("Next gate and reassessment trigger", ""),
        ],
        [3000, 6360],
        font_size=9,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Hard truth",
        "A project can be green while the transformation is red. Delivery health and outcome readiness must be reported separately.",
        fill=GOLD_FILL,
        label_color=GOLD,
    )

    start_page(
        doc,
        "Tool 04 · Cutover",
        "Coexistence, cutover and benefits",
        "The uncomfortable period when old and new systems coexist is a control state in its own right.",
    )
    add_table(
        doc,
        ["CONTROL AREA", "QUESTION", "PASS CRITERION / EVIDENCE", "OWNER"],
        [
            ("Population", "Do both systems process the same in-scope population?", "", ""),
            ("Decision", "How are different outcomes compared and resolved?", "", ""),
            ("Evidence", "Can historical and new decisions be reproduced?", "", ""),
            ("Data", "Are transformations, timing and missingness reconciled?", "", ""),
            ("Operations", "Are dual queues, backlogs and handoffs controlled?", "", ""),
            ("Customer", "Are duplicate requests, restrictions or inconsistent messages prevented?", "", ""),
            ("Rollback", "What can be reversed, how quickly and with what data?", "", ""),
            ("Retirement", "What evidence permits the legacy control to be decommissioned?", "", ""),
        ],
        [1700, 3200, 3100, 1360],
        font_size=8.15,
        first_col_bold=True,
    )
    doc.add_heading("Benefits realization", level=2)
    add_table(
        doc,
        ["OUTCOME", "BASELINE", "TARGET", "CURRENT", "EVIDENCE", "OWNER"],
        [
            ("Risk coverage", "", "", "", "", ""),
            ("Decision quality", "", "", "", "", ""),
            ("Customer friction", "", "", "", "", ""),
            ("Time to adapt", "", "", "", "", ""),
            ("Operational capacity", "", "", "", "", ""),
            ("Control cost", "", "", "", "", ""),
        ],
        [1800, 1150, 1150, 1150, 2600, 1510],
        font_size=8.1,
        header_fill=BLUE_FILL,
        first_col_bold=True,
    )
    add_decision_record(doc, [
        "Cutover decision", "Residual risk and accountable executive",
        "Legacy retirement / benefits review date", "Stop / rollback trigger"
    ])

    # TOOL 05
    add_tool_header(
        doc,
        5,
        "Global Compliance Operating Model",
        "Place standards, capabilities and judgment where they produce consistency without creating a global queue.",
        "Designing global hubs, follow-the-sun operations, regional growth or correcting excessive centralization.",
        "A capability placement map, decision-right matrix, handoff standard and concentration response.",
    )
    add_callout(
        doc,
        "Design principle",
        "Centralize the standard. Distribute informed judgment. Connect the system through evidence.",
    )
    doc.add_heading("Capability placement test", level=2)
    add_table(
        doc,
        ["CAPABILITY", "GLOBAL", "REGIONAL", "LOCAL", "HYBRID", "RATIONALE / EVIDENCE"],
        [
            ("Policy and minimum standards", "", "", "", "", ""),
            ("Regulatory interpretation", "", "", "", "", ""),
            ("Customer / product risk decisions", "", "", "", "", ""),
            ("Investigations and intelligence", "", "", "", "", ""),
            ("Screening / monitoring operations", "", "", "", "", ""),
            ("Quality assurance and calibration", "", "", "", "", ""),
            ("Data, models and control technology", "", "", "", "", ""),
            ("Regulatory engagement and reporting", "", "", "", "", ""),
        ],
        [2200, 800, 900, 800, 800, 3860],
        font_size=8.0,
        header_fill=DARK,
        first_col_bold=True,
    )
    doc.add_heading("Placement criteria", level=2)
    add_table(
        doc,
        ["CENTRALIZE WHEN", "DISTRIBUTE WHEN"],
        [
            ("Consistency, scarce expertise, control integrity or scale materially improve.", "Language, market context, regulatory proximity or response speed materially improve."),
            ("Fragmentation would create incompatible data, models or evidence.", "Local judgment can operate within a clear global standard and escalation rule."),
            ("Independent challenge or enterprise visibility requires separation.", "Customer or regulator outcomes suffer when decisions wait for a remote queue."),
        ],
        [4680, 4680],
        font_size=8.8,
        header_fill=BLUE_FILL,
    )

    start_page(
        doc,
        "Tool 05 · Decision rights",
        "Decision rights and follow-the-sun handoffs",
        "A distributed model works only when authority and evidence travel together.",
    )
    add_table(
        doc,
        ["DECISION", "GLOBAL OWNER", "REGIONAL AUTHORITY", "LOCAL EVIDENCE", "MUST ESCALATE WHEN"],
        [
            ("Policy interpretation", "", "", "", ""),
            ("Customer onboarding / restriction", "", "", "", ""),
            ("Suspicious activity / investigation", "", "", "", ""),
            ("Sanctions or high-risk nexus", "", "", "", ""),
            ("Model / threshold change", "", "", "", ""),
            ("Regulatory communication", "", "", "", ""),
            ("Incident or material control failure", "", "", "", ""),
        ],
        [1900, 1600, 1900, 1900, 2060],
        font_size=8.0,
        header_fill=SOFT,
        first_col_bold=True,
    )
    doc.add_heading("Minimum handoff packet", level=2)
    add_table(
        doc,
        ["HANDOFF ELEMENT", "STANDARD", "OWNER / SYSTEM OF RECORD"],
        [
            ("Decision state", "What has been decided, what remains open and the deadline.", ""),
            ("Evidence", "Relevant customer, transaction, intelligence and prior-decision sources.", ""),
            ("Risk hypothesis", "What concern is being tested and what would disprove it.", ""),
            ("Authority", "What the receiving team may decide and what must escalate.", ""),
            ("Priority", "Consequence and time sensitivity—not merely queue age.", ""),
            ("Closure", "Required rationale, evidence retention and feedback to upstream teams.", ""),
        ],
        [2000, 5000, 2360],
        font_size=8.7,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Failure signal",
        "If the receiving team must reconstruct the case, the organization has not transferred work—it has transferred confusion.",
        fill=RED_FILL,
        label_color=RED,
    )

    start_page(
        doc,
        "Tool 05 · Resilience",
        "Concentration and operating resilience stress test",
        "Test the model against the loss of a location, leader, vendor, dataset or regulatory assumption.",
    )
    add_table(
        doc,
        ["SCENARIO", "CURRENT EXPOSURE", "TOLERANCE", "IMMEDIATE RESPONSE", "RECOVERY EVIDENCE", "OWNER"],
        [
            ("Primary hub unavailable", "", "", "", "", ""),
            ("Key leader / specialist unavailable", "", "", "", "", ""),
            ("Critical vendor or model unavailable", "", "", "", "", ""),
            ("Material data feed degraded", "", "", "", "", ""),
            ("Regulatory interpretation conflicts across markets", "", "", "", "", ""),
            ("Volume surge exceeds regional capacity", "", "", "", "", ""),
            ("Cyber or geopolitical event isolates a location", "", "", "", "", ""),
        ],
        [1950, 1450, 1100, 2000, 1860, 1000],
        font_size=7.8,
        header_fill=DARK,
        first_col_bold=True,
    )
    doc.add_heading("Leadership-depth test", level=2)
    add_table(
        doc,
        ["QUESTION", "EVIDENCE / RESPONSE"],
        [
            ("Can each hub make and defend decisions, or does it only execute?", ""),
            ("Which material decisions still depend on one person, time zone or language?", ""),
            ("Are performance differences visible and converted into shared learning?", ""),
            ("Can regions challenge the global standard without creating local exceptions by default?", ""),
            ("What work could move within 24 hours without losing control evidence?", ""),
        ],
        [4700, 4660],
        font_size=8.8,
        first_col_bold=True,
    )
    add_decision_record(doc, [
        "Operating-model decision", "Risk accepted / accountable executive",
        "Capability to relocate or duplicate", "Target / stress-test date"
    ])
    start_page(
        doc,
        "Tool 05 · Action",
        "Operating-model implementation plan",
        "Turn the placement and resilience decisions into specific authority, capability and evidence changes.",
    )
    add_action_register(doc, rows=6)
    doc.add_heading("Progress review", level=2)
    add_table(
        doc,
        ["REVIEW FIELD", "WORKING RECORD"],
        [
            ("Review cadence and forum", ""),
            ("Calibration / consistency evidence", ""),
            ("Regional or local feedback requiring a standard change", ""),
            ("Decision if concentration remains outside tolerance", ""),
        ],
        [3400, 5960],
        font_size=9,
        first_col_bold=True,
    )

    # TOOL 06
    add_tool_header(
        doc,
        6,
        "Board Oversight of AI Agents",
        "Help directors govern consequence, authority and accountability without turning the board into a technical design committee.",
        "Board or risk-committee approval, oversight of a material AI use case or expansion of agent authority.",
        "A board challenge record, concise dashboard and explicit approval conditions.",
    )
    doc.add_heading("Questions the board should ask", level=2)
    add_table(
        doc,
        ["DOMAIN", "BOARD QUESTION", "EVIDENCE MANAGEMENT SHOULD PROVIDE"],
        [
            ("Purpose", "What outcome is the system designed to improve, and for whom?", "Approved purpose, affected population, baseline and success measure."),
            ("Authority", "What can it do—not merely recommend—and what is the highest-consequence action?", "Authority level, permissions, limits and prohibited actions."),
            ("Human oversight", "Where can a person meaningfully disagree before consequence?", "Workflow, reviewer evidence, override and disagreement data."),
            ("Evidence", "Can an independent reviewer reconstruct the system’s reasoning and actions?", "Version, inputs, outputs, tool calls, approvals and retained rationale."),
            ("Testing", "How was the system tested against realistic failure and adversarial behavior?", "Acceptance, segment, bias, security and adversarial test results."),
            ("Resilience", "How quickly can it be stopped, rolled back or isolated?", "Kill switch, rollback time, incident rehearsal and contingency."),
            ("Change", "Which changes require revalidation or reapproval?", "Change taxonomy covering model, data, tool, purpose, authority and market."),
            ("Third parties", "Where are concentration, continuity and exit risks?", "Dependency map, contract rights, alternative service and exit plan."),
            ("Outcomes", "How will management distinguish productivity from risk moved elsewhere?", "Risk, customer, control, workforce and quality measures."),
            ("Accountability", "Who owns harm when the system follows its instructions?", "Named accountable executive, control owners and incident governance."),
        ],
        [1400, 3950, 4010],
        font_size=8.0,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Board boundary",
        "The board does not need to approve prompts. It should approve the consequence envelope, evidence standard, accountability and conditions for management to expand authority.",
        fill=BLUE_FILL,
        label_color=BLUE,
    )

    start_page(
        doc,
        "Tool 06 · Oversight",
        "Board dashboard and decision record",
        "Keep the board view concise enough to govern, but specific enough to expose false confidence.",
    )
    add_table(
        doc,
        ["DOMAIN", "STATUS", "KEY EVIDENCE", "CHANGE SINCE LAST REVIEW", "MANAGEMENT ACTION", "BOARD CHALLENGE"],
        [
            ("Purpose / population", "", "", "", "", ""),
            ("Authority / permissions", "", "", "", "", ""),
            ("Performance / outcomes", "", "", "", "", ""),
            ("Human oversight", "", "", "", "", ""),
            ("Incidents / near misses", "", "", "", "", ""),
            ("Drift / material change", "", "", "", "", ""),
            ("Third-party / resilience", "", "", "", "", ""),
            ("Residual risk", "", "", "", "", ""),
        ],
        [1500, 850, 1900, 1800, 1800, 1510],
        font_size=7.7,
        header_fill=DARK,
        first_col_bold=True,
    )
    doc.add_heading("Board decision", level=2)
    add_table(
        doc,
        ["FIELD", "BOARD RECORD"],
        [
            ("Decision: approve / conditional / decline / require redesign", ""),
            ("Approved purpose, population and authority ceiling", ""),
            ("Conditions management must satisfy", ""),
            ("Risk accepted and accountable executive", ""),
            ("Metrics / incidents requiring notification", ""),
            ("Changes requiring return to the board", ""),
            ("Next review date", ""),
        ],
        [3500, 5860],
        font_size=9,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "The one question",
        "What would have to be true for management to recommend stopping this system—and would the board hear about it in time?",
        fill=GOLD_FILL,
        label_color=GOLD,
    )

    # TOOL 07
    add_tool_header(
        doc,
        7,
        "30/60/90-Day Transformation Plan",
        "Create visible progress in one quarter without confusing urgency with ungoverned activity.",
        "Starting, resetting or rescuing a compliance transformation.",
        "A one-quarter charter, sequenced evidence plan, ownership model and scale / stop decision.",
    )
    doc.add_heading("Define the 90-day outcome", level=2)
    add_table(
        doc,
        ["CHARTER FIELD", "WORKING DEFINITION"],
        [
            ("Outcome to improve", ""),
            ("Affected risk, customer and operating population", ""),
            ("Baseline and evidence date", ""),
            ("90-day target and acceptable tolerance", ""),
            ("Accountable executive", ""),
            ("Permanent operating owner", ""),
            ("Decision forum and cadence", ""),
            ("What is explicitly out of scope", ""),
            ("Stop condition", ""),
        ],
        [3100, 6260],
        font_size=9,
        first_col_bold=True,
    )
    doc.add_heading("Workstream architecture", level=2)
    add_table(
        doc,
        ["WORKSTREAM", "90-DAY QUESTION", "LEAD"],
        [
            ("Risk and control", "What control outcome changes, and how will coverage be proven?", ""),
            ("Data and technology", "What data, model, platform or integration must become reliable?", ""),
            ("Process and customer", "What friction, handoff, exception or delay will be redesigned?", ""),
            ("People and ownership", "Who will operate, challenge and improve the capability?", ""),
            ("Governance and evidence", "What decisions, thresholds and records make change accountable?", ""),
        ],
        [2200, 5600, 1560],
        font_size=8.7,
        header_fill=BLUE_FILL,
        first_col_bold=True,
    )

    start_page(
        doc,
        "Tool 07 · Days 1–30",
        "Establish the truth",
        "The first month should reduce ambiguity: one outcome, one baseline, named ownership and the real constraints.",
    )
    add_table(
        doc,
        ["DELIVERABLE", "EXECUTIVE STANDARD", "OWNER", "DUE", "EVIDENCE / DECISION"],
        [
            ("Outcome charter", "Problem, affected population, baseline, target and non-goals approved.", "", "", ""),
            ("Current-state control map", "Material decisions, controls, handoffs, exceptions and workarounds visible.", "", "", ""),
            ("Data reality", "Quality, lineage, access, ownership and representative-data gaps documented.", "", "", ""),
            ("Risk and dependency map", "People, vendors, models, systems and regulatory dependencies named.", "", "", ""),
            ("Operating baseline", "Coverage, quality, aging, customer friction, capacity and cost measured.", "", "", ""),
            ("Governance", "Sponsor, permanent owner, decision forum, cadence and escalation rule active.", "", "", ""),
            ("90-day proof point", "One narrow outcome can be changed and evidenced within the quarter.", "", "", ""),
        ],
        [2200, 3900, 1100, 800, 1360],
        font_size=8.0,
        header_fill=DARK,
        first_col_bold=True,
    )
    doc.add_heading("Day-30 decision", level=2)
    add_table(
        doc,
        ["QUESTION", "WORKING RESPONSE"],
        [
            ("Is the original problem statement still valid?", ""),
            ("What assumption failed first?", ""),
            ("Which dependency or evidence gap threatens the 90-day outcome?", ""),
            ("What will be stopped, narrowed or deferred?", ""),
            ("Decision: proceed / reframe / stop", ""),
        ],
        [4100, 5260],
        font_size=9,
        first_col_bold=True,
    )
    add_action_register(doc, rows=5)

    start_page(
        doc,
        "Tool 07 · Days 31–60",
        "Prove the change",
        "The second month should test the target process and control with representative users, data and failure conditions.",
    )
    add_table(
        doc,
        ["DELIVERABLE", "EXECUTIVE STANDARD", "OWNER", "DUE", "EVIDENCE / DECISION"],
        [
            ("Target process", "Legacy steps removed; decisions, handoffs and exceptions designed end to end.", "", "", ""),
            ("Control design", "Control purpose, owner, evidence, threshold and escalation are explicit.", "", "", ""),
            ("Representative test", "Realistic populations, edge cases and material segments are included.", "", "", ""),
            ("Failure / adversarial test", "The program tests how data, models, users and dependencies fail.", "", "", ""),
            ("Coexistence and rollback", "Old/new reconciliation, cutover criteria and recovery are rehearsed.", "", "", ""),
            ("Operating ownership", "Permanent team can operate, challenge and improve the capability.", "", "", ""),
            ("Outcome evidence", "Early results compare with baseline and expose unintended effects.", "", "", ""),
        ],
        [2200, 3900, 1100, 800, 1360],
        font_size=8.0,
        header_fill=DARK,
        first_col_bold=True,
    )
    doc.add_heading("Day-60 decision", level=2)
    add_table(
        doc,
        ["QUESTION", "WORKING RESPONSE"],
        [
            ("What has been proven with representative evidence?", ""),
            ("What remains assertion or demonstration?", ""),
            ("What control, customer or operational effect was unintended?", ""),
            ("What must be true before limited release or cutover?", ""),
            ("Decision: release narrowly / continue proving / redesign / stop", ""),
        ],
        [4200, 5160],
        font_size=9,
        first_col_bold=True,
    )
    add_action_register(doc, rows=5)

    start_page(
        doc,
        "Tool 07 · Days 61–90",
        "Govern the scale",
        "The final month should transfer ownership, prove sustained outcomes and make the next scale decision explicit.",
    )
    add_table(
        doc,
        ["DELIVERABLE", "EXECUTIVE STANDARD", "OWNER", "DUE", "EVIDENCE / DECISION"],
        [
            ("Outcome comparison", "Current results compared with baseline, target and segment tolerances.", "", "", ""),
            ("Control assurance", "Quality, validation and independent challenge support continued use.", "", "", ""),
            ("Permanent ownership", "Service, model, process and policy owners accept accountability.", "", "", ""),
            ("Monitoring and triggers", "Dashboard, thresholds, escalation and reassessment events are live.", "", "", ""),
            ("Residual risk", "Known gaps, compensating controls, acceptance and expiry dates recorded.", "", "", ""),
            ("Scale economics", "Capacity, cost, customer and risk impacts support the next investment.", "", "", ""),
            ("Learning record", "What changed, what failed and what should not be repeated are published.", "", "", ""),
        ],
        [2200, 3900, 1100, 800, 1360],
        font_size=8.0,
        header_fill=DARK,
        first_col_bold=True,
    )
    add_decision_record(doc, [
        "90-day outcome achieved / missed", "Scale decision", "Evidence supporting the decision",
        "Residual risk", "Permanent accountable owner", "Next review date", "Stop / reassessment trigger"
    ])
    add_callout(
        doc,
        "Close the quarter",
        "Publish the learning—not only the success. An adaptive organization becomes faster because it retains evidence from what did not work.",
        fill=BLUE_FILL,
        label_color=BLUE,
    )

    # CONSOLIDATED OUTPUTS
    start_page(
        doc,
        "Consolidated outputs",
        "Executive action and evidence register",
        "Use this page to bring decisions from multiple tools into one accountable view.",
    )
    add_table(
        doc,
        ["TOOL / DECISION", "ACTION", "OWNER", "DUE", "SUCCESS MEASURE", "STATUS", "TRIGGER"],
        [("", "", "", "", "", "", "") for _ in range(10)],
        [1500, 2100, 1100, 800, 1800, 900, 1160],
        font_size=7.7,
        header_fill=DARK,
    )
    doc.add_heading("Evidence inventory", level=2)
    add_table(
        doc,
        ["EVIDENCE ITEM", "SOURCE / SYSTEM", "OWNER", "AS-OF DATE", "CONF.", "GAP / ACTION"],
        [("", "", "", "", "", "") for _ in range(8)],
        [2100, 1950, 1250, 1050, 1050, 1960],
        font_size=8.0,
        header_fill=SOFT,
    )
    add_callout(
        doc,
        "Closure standard",
        "An action is not closed because work was performed. It is closed when the agreed evidence demonstrates the intended outcome and the residual risk is understood.",
        fill=GOLD_FILL,
        label_color=GOLD,
    )

    start_page(
        doc,
        "Facilitator close",
        "The tool is the conversation it makes possible",
        "Adaptive compliance is the organizational capability to see change, interpret evidence, decide accountably and improve while risk is still moving.",
    )
    add_table(
        doc,
        ["BEFORE THE SESSION", "IN THE ROOM", "AFTER THE DECISION"],
        [
            (
                "Define scope and decision.\nSelect participants who own evidence and consequence.\nCirculate baseline material.\nCollect independent scores.",
                "Separate maturity from confidence.\nAsk what evidence would change the answer.\nRecord dissent and uncertainty.\nName the decision—not merely the discussion.",
                "Assign owner and due date.\nRetain the evidence reviewed.\nSet success measure and trigger.\nReturn when the trigger occurs, not only on the calendar.",
            )
        ],
        [3120, 3120, 3120],
        font_size=9.2,
        header_fill=DARK,
    )
    add_callout(
        doc,
        "Final principle",
        "Better compliance begins when the organization can change its mind for the right reason—and show its work.",
    )
    doc.add_heading("About the author", level=2)
    doc.add_paragraph(
        "Micheal Sheehy is a global payments and compliance executive and the author of the eight-volume "
        "Adaptive Compliance series. His work focuses on financial crime, payment transparency, artificial "
        "intelligence, model governance, global operating models and the leadership systems required to scale safely."
    )
    add_table(
        doc,
        ["READ", "CONTACT"],
        [
            ("michealsheehy.com/adaptive-compliance.html", "micheal@michealsheehy.com  ·  +1 (818) 397-8890"),
            ("michealsheehy.com/tools.html", "linkedin.com/in/micheal-s-88479045"),
        ],
        [4680, 4680],
        font_size=9.5,
        header_fill=BLUE_FILL,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        p.add_run("Vendor-neutral working material · Not legal advice · Confirm applicable regulatory requirements"),
        size=8.5,
        color=MUTED,
        italic=True,
    )

    path = OUTPUT_DIR / "Adaptive_Compliance_Practical_Toolkit.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_toolkit())
